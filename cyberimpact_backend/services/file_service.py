# cyberimpact_backend/services/file_service.py
import os
from pathlib import Path
from typing import Dict, Any, Optional
import openpyxl
from docx import Document
import PyPDF2
from config import UPLOAD_DIR, MAX_FILE_SIZE, ALLOWED_ASSET_EXTENSIONS, ALLOWED_FINANCIAL_EXTENSIONS

class FileService:
    
    @staticmethod
    def ensure_upload_dir():
        """Ensure upload directory exists"""
        os.makedirs(UPLOAD_DIR, exist_ok=True)
    
    @staticmethod
    def validate_file_size(file_size: int) -> bool:
        """Validate file size"""
        return file_size <= MAX_FILE_SIZE
    
    @staticmethod
    def validate_asset_file(filename: str) -> bool:
        """Validate asset inventory file extension"""
        ext = Path(filename).suffix.lower()
        return ext in ALLOWED_ASSET_EXTENSIONS
    
    @staticmethod
    def validate_financial_file(filename: str) -> bool:
        """Validate financial document file extension"""
        ext = Path(filename).suffix.lower()
        return ext in ALLOWED_FINANCIAL_EXTENSIONS
    
    @staticmethod
    def process_xlsx_file(file_path: str) -> Dict[str, Any]:
        """Process Excel file and extract data"""
        try:
            workbook = openpyxl.load_workbook(file_path, data_only=True)
            data = {
                "sheets": [],
                "total_sheets": len(workbook.sheetnames)
            }
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                sheet_data = {
                    "name": sheet_name,
                    "rows": [],
                    "max_row": sheet.max_row,
                    "max_column": sheet.max_column
                }
                
                # Extract data (limit to first 1000 rows to avoid memory issues)
                max_rows = min(sheet.max_row, 1000)
                for row in sheet.iter_rows(min_row=1, max_row=max_rows, values_only=True):
                    # Convert row to list, handling None values
                    row_data = [str(cell) if cell is not None else "" for cell in row]
                    sheet_data["rows"].append(row_data)
                
                data["sheets"].append(sheet_data)
            
            return data
        except Exception as e:
            raise ValueError(f"Failed to process Excel file: {str(e)}")
    
    @staticmethod
    def process_pdf_file(file_path: str) -> Dict[str, Any]:
        """Process PDF file and extract text"""
        try:
            extracted_text = ""
            metadata = {}
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                metadata["page_count"] = len(pdf_reader.pages)
                
                # Extract text from all pages
                for page_num, page in enumerate(pdf_reader.pages):
                    text = page.extract_text()
                    extracted_text += f"\n--- Page {page_num + 1} ---\n{text}"
                
                # Get PDF metadata
                if pdf_reader.metadata:
                    metadata["title"] = pdf_reader.metadata.get('/Title', '')
                    metadata["author"] = pdf_reader.metadata.get('/Author', '')
                    metadata["subject"] = pdf_reader.metadata.get('/Subject', '')
            
            return {
                "extracted_text": extracted_text,
                "metadata": metadata
            }
        except Exception as e:
            raise ValueError(f"Failed to process PDF file: {str(e)}")
    
    @staticmethod
    def process_docx_file(file_path: str) -> Dict[str, Any]:
        """Process Word document and extract text"""
        try:
            doc = Document(file_path)
            extracted_text = ""
            
            # Extract text from paragraphs
            for para in doc.paragraphs:
                extracted_text += para.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        extracted_text += cell.text + "\t"
                    extracted_text += "\n"
            
            metadata = {
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables)
            }
            
            return {
                "extracted_text": extracted_text,
                "metadata": metadata
            }
        except Exception as e:
            raise ValueError(f"Failed to process Word document: {str(e)}")
    
    @staticmethod
    def save_uploaded_file(file_content: bytes, filename: str) -> str:
        """Save uploaded file to disk"""
        FileService.ensure_upload_dir()
        file_path = os.path.join(UPLOAD_DIR, filename)
        
        with open(file_path, 'wb') as f:
            f.write(file_content)
        
        return file_path
    
    @staticmethod
    def delete_file(file_path: str):
        """Delete file from disk"""
        try:
            if os.path.exists(file_path):
                os.remove(file_path)
        except Exception as e:
            print(f"Failed to delete file {file_path}: {e}")
